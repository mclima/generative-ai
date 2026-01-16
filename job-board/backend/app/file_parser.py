from typing import Optional
import io
from PyPDF2 import PdfReader
from docx import Document
import zipfile

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        
        result = "\n".join(text)
        if not result.strip():
            raise ValueError("PDF appears to be empty or contains only images")
        return result
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        # First check if it's a valid zip file (DOCX is a zip archive)
        docx_file = io.BytesIO(file_content)
        if not zipfile.is_zipfile(docx_file):
            raise ValueError("File is not a valid DOCX format. Please ensure you're uploading a .docx file (not .doc). You can convert .doc files to .docx in Microsoft Word or use 'Save As' and select .docx format.")
        
        # Reset the file pointer
        docx_file.seek(0)
        doc = Document(docx_file)
        
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        result = "\n".join(text)
        if not result.strip():
            raise ValueError("DOCX file appears to be empty")
        return result
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

def parse_resume_file(filename: str, file_content: bytes) -> str:
    """Parse resume file and extract text based on file type"""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_content)
    elif filename_lower.endswith('.doc'):
        raise ValueError("Old .doc format is not supported. Please convert to .docx format using Microsoft Word or Google Docs.")
    elif filename_lower.endswith('.txt'):
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            return file_content.decode('latin-1')
    else:
        raise ValueError(f"Unsupported file type. Please upload PDF, DOCX, or TXT files.")
